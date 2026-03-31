#!/usr/bin/env python3
"""Create D&D monster stat cards, 2 per letter-sized page"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    """Add a horizontal line after a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '8B0000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_monster_card(doc, monster_data):
    """Create a single monster card"""
    
    # Monster name (title)
    title = doc.add_paragraph()
    title_run = title.add_run(monster_data['name'].upper())
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(title)
    
    # Type/size info
    if 'type' in monster_data:
        type_para = doc.add_paragraph(monster_data['type'])
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        type_para_format = type_para.paragraph_format
        type_para_format.space_after = Pt(6)
    
    # Hit Points
    hp_para = doc.add_paragraph()
    hp_run = hp_para.add_run('Hit Points: ')
    hp_run.font.bold = True
    hp_para.add_run(str(monster_data['hp']))
    
    # Armor Class
    ac_para = doc.add_paragraph()
    ac_run = ac_para.add_run('Armor Class: ')
    ac_run.font.bold = True
    ac_para.add_run(str(monster_data['ac']))
    
    # Spacing
    doc.add_paragraph()
    
    # Special Abilities
    if 'abilities' in monster_data and monster_data['abilities']:
        abilities_header = doc.add_paragraph()
        abilities_header_run = abilities_header.add_run('SPECIAL ABILITIES')
        abilities_header_run.font.bold = True
        abilities_header_run.font.size = Pt(11)
        abilities_header_run.font.color.rgb = RGBColor(139, 0, 0)
        
        for ability in monster_data['abilities']:
            ability_para = doc.add_paragraph(style='List Bullet')
            ability_para.add_run(ability)
    
    # Attacks
    if 'attacks' in monster_data and monster_data['attacks']:
        attacks_header = doc.add_paragraph()
        attacks_header_run = attacks_header.add_run('ATTACKS')
        attacks_header_run.font.bold = True
        attacks_header_run.font.size = Pt(11)
        attacks_header_run.font.color.rgb = RGBColor(139, 0, 0)
        
        for attack in monster_data['attacks']:
            attack_para = doc.add_paragraph(style='List Bullet')
            attack_para.add_run(attack)

# Monster data
monsters = [
    {
        'name': 'Orc',
        'type': 'Medium humanoid (orc), chaotic evil',
        'hp': 15,
        'ac': 13,
        'abilities': [],
        'attacks': [
            'Greataxe: Melee Weapon Attack, 1d12 + 3 slashing damage',
            'Javelin: Melee or Ranged Weapon Attack, 1d6 + 3 piercing damage'
        ]
    },
    {
        'name': 'Wight',
        'type': 'Medium undead, neutral evil',
        'hp': 45,
        'ac': 14,
        'abilities': [
            'Life Drain: Melee Weapon Attack, target must succeed on DC 13 Constitution saving throw or its hit point maximum is reduced by an amount equal to the damage taken'
        ],
        'attacks': [
            'Longsword: Melee Weapon Attack, 1d8 + 2 slashing damage (or 1d10 + 2 if two-handed)',
            'Longbow: Ranged Weapon Attack, 1d8 + 2 piercing damage'
        ]
    },
    {
        'name': 'Specter',
        'type': 'Medium undead, chaotic evil',
        'hp': 22,
        'ac': 12,
        'abilities': [
            'Incorporeal Movement: The specter can move through other creatures and objects as if they were difficult terrain. It takes 5 (1d10) force damage if it ends its turn inside an object',
            'Life Drain: Melee Spell Attack, 3d6 necrotic damage. Target must succeed on DC 10 Constitution saving throw or its hit point maximum is reduced by amount equal to damage taken'
        ],
        'attacks': []
    },
    {
        'name': 'Skeleton',
        'type': 'Medium undead, lawful evil',
        'hp': 13,
        'ac': 13,
        'abilities': [],
        'attacks': [
            'Shortsword: Melee Weapon Attack, 1d6 + 2 piercing damage',
            'Shortbow: Ranged Weapon Attack, 1d6 + 2 piercing damage'
        ]
    },
    {
        'name': 'Giant Spider',
        'type': 'Large beast, unaligned',
        'hp': 26,
        'ac': 14,
        'abilities': [
            'Web (Recharge 5-6): Ranged Weapon Attack, target is restrained by webbing. DC 12 Strength check to escape. The webbing can be attacked (AC 10, 5 HP)'
        ],
        'attacks': [
            'Bite: Melee Weapon Attack, 1d8 + 3 piercing damage plus 2d8 poison damage. Target must succeed on DC 11 Constitution saving throw or take poison damage (half on success)'
        ]
    },
    {
        'name': 'Wolf Spider',
        'type': 'Medium beast, unaligned',
        'hp': 11,
        'ac': 13,
        'abilities': [
            'Spider Climb: The spider can climb difficult surfaces, including upside down on ceilings, without needing to make an ability check',
            'Web Sense: While in contact with a web, the spider knows the exact location of any creature in contact with the same web',
            'Web Walker: The spider ignores movement restrictions caused by webbing'
        ],
        'attacks': [
            'Bite: Melee Weapon Attack, 1d6 + 1 piercing damage plus 1d6 poison damage. Target must succeed on DC 10 Constitution saving throw or take poison damage (half on success)'
        ]
    },
    {
        'name': 'Ettercap',
        'type': 'Medium monstrosity, neutral evil',
        'hp': 44,
        'ac': 13,
        'abilities': [
            'Spider Climb: The ettercap can climb difficult surfaces, including upside down on ceilings, without needing to make an ability check',
            'Web Sense: While in contact with a web, the ettercap knows the exact location of any creature in contact with the same web',
            'Web Walker: The ettercap ignores movement restrictions caused by webbing',
            'Web (Recharge 5-6): Ranged Weapon Attack, target is restrained by webbing. DC 11 Strength check to escape'
        ],
        'attacks': [
            'Bite: Melee Weapon Attack, 1d8 + 2 piercing damage plus 1d8 poison damage. Target must succeed on DC 11 Constitution saving throw or be poisoned for 1 minute',
            'Claws: Melee Weapon Attack, 2d4 + 2 slashing damage'
        ]
    }
]

# Create document
doc = Document()

# Set margins for half-page cards
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Create cards (2 per page)
for i, monster in enumerate(monsters):
    create_monster_card(doc, monster)
    
    # Add page break after every other card (except the last one)
    if i % 2 == 1 and i < len(monsters) - 1:
        doc.add_page_break()
    elif i % 2 == 0 and i < len(monsters) - 1:
        # Add spacing between cards on same page
        doc.add_paragraph()
        separator = doc.add_paragraph('─' * 60)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

# Save document
output_path = '/Users/moon/MoveAlong/dnd_monster_cards.docx'
doc.save(output_path)
print(f"Monster cards created successfully at {output_path}")
